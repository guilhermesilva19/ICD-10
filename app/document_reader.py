"""Document text extraction utilities"""

import io
from typing import Optional


def extract_text_from_file(file_content: bytes, filename: str) -> Optional[str]:
    """
    Extract text from uploaded document
    
    Args:
        file_content: Raw file bytes
        filename: Original filename with extension
        
    Returns:
        str: Extracted text or None if unsupported format
    """
    file_extension = filename.lower().split('.')[-1]
    
    try:
        if file_extension == 'txt':
            return _extract_from_txt(file_content)
        elif file_extension == 'pdf':
            return _extract_from_pdf(file_content)
        elif file_extension in ['doc', 'docx']:
            return _extract_from_docx(file_content)
        elif file_extension in ['html', 'htm']:
            return _extract_from_html(file_content)
        else:
            return None
    except Exception as e:
        print(f"Error extracting text from {filename}: {str(e)}")
        return None


def extract_title_from_file(file_content: bytes, filename: str) -> Optional[str]:
    """
    Extract title from filename by removing date and extension
    
    Expected formats:
    - PDF: "Macular Degeneration, Age-Related 05-24-2025.pdf" → "Macular Degeneration, Age-Related"
    - HTML: "Allergens, Bedroom Dust Mites.html" → "Allergens, Bedroom Dust Mites"
    
    Args:
        file_content: Raw file bytes (unused but kept for compatibility)
        filename: Original filename with extension
        
    Returns:
        str: Clean title without date and extension
    """
    try:
        # Remove file extension
        name_without_ext = filename.rsplit('.', 1)[0]
        
        # Remove date pattern (MM-DD-YYYY) from the end
        import re
        # Pattern matches: space + MM-DD-YYYY at the end
        clean_title = re.sub(r'\s+\d{2}-\d{2}-\d{4}$', '', name_without_ext)
        
        return clean_title.strip() if clean_title else name_without_ext
        
    except Exception as e:
        print(f"Error extracting title from filename {filename}: {str(e)}")
        return filename.rsplit('.', 1)[0]  # Fallback: just remove extension





def _extract_from_txt(file_content: bytes) -> str:
    """Extract text from plain text file"""
    try:
        return file_content.decode('utf-8')
    except UnicodeDecodeError:
        return file_content.decode('latin-1')


def _extract_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file"""
    try:
        import PyPDF2
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        
        return text.strip()
    except ImportError:
        raise Exception("PyPDF2 not installed. Install with: pip install PyPDF2")





def _extract_from_docx(file_content: bytes) -> str:
    """Extract text from Word document"""
    try:
        from docx import Document
        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)
        
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        return text.strip()
    except ImportError:
        raise Exception("python-docx not installed. Install with: pip install python-docx")


def _extract_from_html(file_content: bytes) -> str:
    """Extract text from HTML file"""
    try:
        from bs4 import BeautifulSoup
        
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                html_content = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            return ""
        
        soup = BeautifulSoup(html_content, 'html.parser')
        return soup.get_text().strip()
        
    except ImportError:
        raise Exception("BeautifulSoup not installed. Install with: pip install beautifulsoup4")


def extract_first_page_content(file_content: bytes, filename: str, max_chars: int = 500) -> str:
    """
    Extract first page content for enhanced vector search
    
    Args:
        file_content: Raw file bytes
        filename: Original filename with extension
        max_chars: Maximum characters to extract (default 500)
        
    Returns:
        str: First page content or empty string if extraction fails
    """
    file_extension = filename.lower().split('.')[-1]
    
    try:
        if file_extension == 'pdf':
            return _extract_first_page_from_pdf(file_content, max_chars)
        elif file_extension in ['html', 'htm']:
            return _extract_first_page_from_html(file_content, max_chars)
        elif file_extension in ['doc', 'docx']:
            return _extract_first_page_from_docx(file_content, max_chars)
        elif file_extension == 'txt':
            return _extract_first_page_from_txt(file_content, max_chars)
        else:
            print(f"⚠️ Unsupported file type for first page extraction: {file_extension}")
            return ""
    except Exception as e:
        print(f"❌ First page extraction failed for {filename}: {str(e)}")
        return ""


def _extract_first_page_from_pdf(file_content: bytes, max_chars: int) -> str:
    """Extract first page from PDF file"""
    try:
        import PyPDF2
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        if len(pdf_reader.pages) == 0:
            print("⚠️ PDF has no pages")
            return ""
        
        # Extract ONLY first page (page 0)
        first_page_text = pdf_reader.pages[0].extract_text()
        
        # Limit to max_chars
        limited_text = first_page_text[:max_chars] if first_page_text else ""
        
        print(f"📄 PDF First Page Extracted: {len(limited_text)} chars")
        print(f"🔍 First Page Preview: {limited_text[:100]}...")
        
        return limited_text.strip()
        
    except ImportError:
        print("❌ PyPDF2 not available for PDF first page extraction")
        return ""


def _extract_first_page_from_html(file_content: bytes, max_chars: int) -> str:
    """Extract first section from HTML file"""
    try:
        from bs4 import BeautifulSoup
        
        # Try different encodings
        for encoding in ['utf-8', 'latin-1', 'cp1252']:
            try:
                html_content = file_content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            print("⚠️ Could not decode HTML content")
            return ""
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Get first meaningful content section
        # content_text = soup.body.get_text()
        # limited_text = content_text[:max_chars] if content_text else ""
        
        body = soup.body

        output_lines = []
        heading_seen = False

        for tag in body.children:
            if getattr(tag, 'name', None) is None:
                continue

            if tag.name in ['h1', 'h2', 'h3', 'h4']:
                if heading_seen:
                    break
                heading_seen = True

            if tag.name not in ['img', 'script', 'style']:
                text = tag.get_text(strip=True)
                if text:
                    output_lines.append(text)

        limited_text = '\n'.join(output_lines)

        print(f"🌐 HTML First Section Extracted: {len(limited_text)} chars")
        print(f"🔍 First Section Preview: {limited_text}...")
        
        return limited_text.strip()
        
    except ImportError:
        print("❌ BeautifulSoup not available for HTML first page extraction")
        return ""


def _extract_first_page_from_docx(file_content: bytes, max_chars: int) -> str:
    """Extract first section from Word document"""
    try:
        from docx import Document
        doc_file = io.BytesIO(file_content)
        doc = Document(doc_file)
        
        # Extract text until we reach max_chars (approximate first page)
        text_parts = []
        current_length = 0
        
        for paragraph in doc.paragraphs:
            if current_length >= max_chars:
                break
            para_text = paragraph.text
            if current_length + len(para_text) > max_chars:
                # Take partial paragraph to reach max_chars
                remaining = max_chars - current_length
                text_parts.append(para_text[:remaining])
                break
            text_parts.append(para_text)
            current_length += len(para_text)
        
        limited_text = " ".join(text_parts)
        
        print(f"📝 DOCX First Section Extracted: {len(limited_text)} chars")
        print(f"🔍 First Section Preview: {limited_text[:100]}...")
        
        return limited_text.strip()
        
    except ImportError:
        print("❌ python-docx not available for DOCX first page extraction")
        return ""


def _extract_first_page_from_txt(file_content: bytes, max_chars: int) -> str:
    """Extract first section from text file"""
    try:
        try:
            text_content = file_content.decode('utf-8')
        except UnicodeDecodeError:
            text_content = file_content.decode('latin-1')
        
        limited_text = text_content[:max_chars]
        
        print(f"📄 TXT First Section Extracted: {len(limited_text)} chars")
        print(f"🔍 First Section Preview: {limited_text[:100]}...")
        
        return limited_text.strip()
        
    except Exception as e:
        print(f"❌ TXT first page extraction failed: {e}")
        return ""


 